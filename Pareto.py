import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import PatternFill
import math
import os
import matplotlib.ticker as mtick

st.set_page_config(page_title="Analyse Pareto", page_icon="📊", layout="wide")
st.title("📊 Analyse Pareto – Causes du problème")

# --- Zone pour ajouter un Titre de Graphe ---
graph_title = st.sidebar.text_input("Titre du graphe", "Analyse Pareto")

# --- Base de données interne ---
if "data" not in st.session_state:
    st.session_state.data = pd.DataFrame(columns=["Cause", "Occurrence"])

st.sidebar.header("Ajouter / Supprimer / Importer")

# ------------------------------------------------------------
# 🔥 Importation de fichier
# ------------------------------------------------------------
uploaded_file = st.sidebar.file_uploader(
    "Importer un fichier (CSV ou Excel) avec Cause + Occurrence",
    type=["csv", "xlsx"]
)

if uploaded_file is not None:
    try:
        if uploaded_file.name.endswith(".csv"):
            df_import = pd.read_csv(uploaded_file)
        else:
            df_import = pd.read_excel(uploaded_file)

        # Vérifier colonnes
        if not {"Cause", "Occurrence"}.issubset(df_import.columns):
            st.sidebar.error("Le fichier doit contenir les colonnes : 'Cause' et 'Occurrence'")
        else:
            df_import["Cause"] = df_import["Cause"].astype(str).str.strip()
            df_import["Occurrence"] = df_import["Occurrence"].astype(float)

            existing = st.session_state.data["Cause"].str.lower().tolist()
            df_import = df_import[~df_import["Cause"].str.lower().isin(existing)]

            st.session_state.data = pd.concat([st.session_state.data, df_import], ignore_index=True)
            st.sidebar.success("Importation réussie !")
    except Exception as e:
        st.sidebar.error(f"Erreur d’importation : {e}")

# --- Ajouter une cause ---
with st.sidebar.form("add_form"):
    cause_in = st.text_input("Nom de la cause")
    occ_in = st.number_input("Occurrence", min_value=0.0, step=0.1, format="%.2f")
    submit = st.form_submit_button("Ajouter")

if submit:
    cause_clean = cause_in.strip()
    if cause_clean == "":
        st.sidebar.warning("Nom de cause vide — saisie ignorée.")
    else:
        existing = st.session_state.data["Cause"].str.lower().tolist()
        if cause_clean.lower() in existing:
            st.sidebar.warning("Cette cause existe déjà.")
        else:
            st.session_state.data.loc[len(st.session_state.data)] = [cause_clean, float(occ_in)]
            st.sidebar.success("Cause ajoutée !")

# --- Supprimer une cause ---
unique_causes = st.session_state.data["Cause"].unique().tolist()
delete_choice = st.sidebar.selectbox("Supprimer une cause", [""] + unique_causes)
if st.sidebar.button("Supprimer") and delete_choice != "":
    st.session_state.data = st.session_state.data[st.session_state.data["Cause"] != delete_choice]
    st.sidebar.success(f"Cause '{delete_choice}' supprimée.")

# --- Si vide ---
if st.session_state.data.empty:
    st.warning("Aucune cause ajoutée.")
    st.stop()

# --- TRI DÉCROISSANT ---
df = st.session_state.data.copy()
df["Occurrence"] = df["Occurrence"].astype(float)
df = df.sort_values(by="Occurrence", ascending=False).reset_index(drop=True)

# --- Calcul Pareto ---
df["%"] = df["Occurrence"] / df["Occurrence"].sum()
df["% Cum"] = df["%"].cumsum()

st.subheader("📋 Tableau trié par occurrence décroissante")
st.dataframe(df.style.format({"Occurrence": "{:.3f}", "%": "{:.2%}", "% Cum": "{:.2%}"}))

# --- Fonction pour convertir figure -> PNG ---
def fig_to_png_bytes(figure):
    buf = BytesIO()
    figure.savefig(buf, format="png", bbox_inches='tight')
    buf.seek(0)
    return buf

# --- Calcul Top causes (cumul ≤ 80%) ---
cum_pct = df["% Cum"]
top_count = cum_pct[cum_pct <= 0.8].index.max() + 1
if math.isnan(top_count) or top_count == 0:
    top_count = 1

# --- Graphe Pareto ---
st.subheader("📉📈 Diagramme de Pareto")

fig, ax1 = plt.subplots(figsize=(10, 5))
plt.title(graph_title)

x = np.arange(len(df))
ax1.bar(x, df["Occurrence"], align='center')
ax1.set_xticks(x)
ax1.set_xticklabels(df["Cause"], rotation=45, ha='right')
ax1.set_ylabel("Occurrence")

ax2 = ax1.twinx()
ax2.plot(x, df["% Cum"], marker="o", linewidth=2)
ax2.yaxis.set_major_formatter(mtick.PercentFormatter(1.0))
ax2.set_ylim(0, 1.05)
ax2.set_ylabel("Pourcentage cumulé (%)")

# Ligne rouge 80%
ax2.axhline(0.8, color='red', linestyle='--', linewidth=1.5)

# Rectangle Top causes
rect = Rectangle((-0.5, 0), top_count, 1.05, facecolor='orange', alpha=0.15)
ax2.add_patch(rect)

fig.tight_layout()
st.pyplot(fig)
png_full = fig_to_png_bytes(fig)

# --- TOP 20% (Top causes cumulant 80%) ---
fig_top, ax1t = plt.subplots(figsize=(8, 4))
plt.title(graph_title + " – Top causes")

df_top = df.iloc[:top_count].copy()
xt = np.arange(len(df_top))

ax1t.bar(xt, df_top["Occurrence"])
ax1t.set_xticks(xt)
ax1t.set_xticklabels(df_top["Cause"], rotation=45, ha='right')

ax2t = ax1t.twinx()
df_top["%"] = df_top["Occurrence"] / df["Occurrence"].sum()
df_top["% Cum (global)"] = df_top["%"].cumsum()
ax2t.plot(xt, df_top["% Cum (global)"], marker='o')
ax2t.set_ylim(0, 1.05)
ax2t.axhline(0.8, color='red', linestyle='--')

rect_top = Rectangle((-0.5, 0), len(df_top), 1.05, facecolor='orange', alpha=0.12)
ax2t.add_patch(rect_top)

fig_top.tight_layout()
png_top = fig_to_png_bytes(fig_top)

# --- EXPORT EXCEL ---
def export_excel_openpyxl(df, img_bytes, title):
    wb = Workbook()
    ws = wb.active
    ws.title = "Pareto"

    ws["A1"] = title
    ws.append(["Cause", "Occurrence", "%", "% Cum"])

    for _, r in df.iterrows():
        ws.append([r["Cause"], float(r["Occurrence"]), float(r["%"]), float(r["% Cum"])])

    img_bytes.seek(0)
    img = XLImage(img_bytes)
    ws.add_image(img, "F3")

    fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
    for row in range(3, 3 + top_count):
        ws[f"A{row}"].fill = fill
        ws[f"B{row}"].fill = fill

    out = BytesIO()
    wb.save(out)
    out.seek(0)
    return out

excel_bytes = export_excel_openpyxl(df, png_full, graph_title)

st.download_button(
    label="📥 Télécharger Excel",
    data=excel_bytes,
    file_name=f"{graph_title}_Pareto_Complet.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)

# --- EXPORT WORD ---
def generate_word_report(df, png_top_bytes, top_count, title):
    tmp_top = "pareto_top_temp.png"

    with open(tmp_top, "wb") as f:
        png_top_bytes.seek(0)
        f.write(png_top_bytes.read())

    doc = Document()
    doc.add_heading(f"Rapport d'Analyse Pareto – {title}", level=1)

    doc.add_heading("1. Résumé", level=2)
    total_occ = df["Occurrence"].sum()
    cover_top_percent = df['% Cum'].iloc[top_count-1] * 100

    doc.add_paragraph(f"- Nombre total de causes : {len(df)}")
    doc.add_paragraph(f"- Causes principales : {top_count}")
    doc.add_paragraph(f"- Elles couvrent {cover_top_percent:.1f}% des occurrences")

    doc.add_heading("2. Graphique Top causes", level=2)
    doc.add_picture(tmp_top, width=Inches(6))

    doc.add_heading("3. Interprétation", level=2)
    doc.add_paragraph(
        "Les causes principales (top causes) doivent être traitées en priorité "
        "car elles génèrent la majorité du problème."
    )

    os.remove(tmp_top)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

word_bytes = generate_word_report(df, png_top, top_count, graph_title)

st.download_button(
    label="📄 Télécharger Rapport Word",
    data=word_bytes,
    file_name=f"{graph_title}_Rapport_Pareto.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

st.info("Le Word contient le Top causes cumulant 80% — L’Excel contient toutes les données + graphe complet.")

# --- CSS pour animations ---
st.markdown("""
<style>
/* Animation d'apparition pour le body */
body {
    animation: fadeIn 1s ease-in;
}

/* Keyframes pour fadeIn */
@keyframes fadeIn {
    from { opacity: 0; }
    to { opacity: 1; }
}

/* Titres H2 et H3 animés au survol */
h2, h3 {
    transition: color 0.3s, transform 0.3s;
}
h2:hover, h3:hover {
    color: #0055aa !important;
    transform: scale(1.05);
}

/* Boutons animés */
.stButton>button {
    transition: transform 0.2s, background-color 0.3s;
}
.stButton>button:hover {
    transform: scale(1.05);
    background-color: #003f7d !important;
}

/* Badge React */
#react-root div {
    transition: transform 0.5s, box-shadow 0.5s;
}
#react-root div:hover {
    transform: scale(1.1);
    box-shadow: 0px 5px 15px rgba(0,0,0,0.3);
}
</style>
""", unsafe_allow_html=True)

# --- JavaScript + React pour Badge animé ---
st.components.v1.html("""
<div id="react-root"></div>

<script src="https://unpkg.com/react@18/umd/react.development.js"></script>
<script src="https://unpkg.com/react-dom@18/umd/react-dom.development.js"></script>

<script>
const root = ReactDOM.createRoot(document.getElementById("react-root"));

function Badge() {
    const [scale, setScale] = React.useState(1);

    return React.createElement(
        "div",
        {
            style: {
                backgroundColor: "#0055aa",
                color: "white",
                padding: "10px 20px",
                borderRadius: "12px",
                fontWeight: "bold",
                display: "inline-block",
                marginBottom: "20px",
                transform: `scale(${scale})`,
                transition: "transform 0.3s, box-shadow 0.3s",
                cursor: "pointer",
                boxShadow: "0px 3px 10px rgba(0,0,0,0.2)"
            },
            onMouseEnter: () => setScale(1.1),
            onMouseLeave: () => setScale(1)
        }, 
}

root.render(React.createElement(Badge));
</script>
""", height=80)
st.components.v1.html("""
<html>
<head>
  <style>
    body, html { margin:0; padding:0; height:100%; overflow:hidden; }
    #bgCanvas { position: fixed; top:0; left:0; width:100%; height:100%; z-index:-1; }
  </style>
</head>
<body>
<canvas id="bgCanvas"></canvas>
<div style="position:relative; z-index:1; text-align:center; margin-top:50px;">
  <h1 style="color:white; text-shadow: 2px 2px 5px black;"> Analyse Pareto – AR Effect</h1>
</div>

<script src="https://cdnjs.cloudflare.com/ajax/libs/three.js/r152/three.min.js"></script>
<script>
  // 1️⃣ Création scène, caméra et renderer
  const scene = new THREE.Scene();
  const camera = new THREE.PerspectiveCamera(75, window.innerWidth/window.innerHeight, 0.1, 1000);
  const renderer = new THREE.WebGLRenderer({canvas: document.getElementById("bgCanvas"), alpha:true});
  renderer.setSize(window.innerWidth, window.innerHeight);

  // 2️⃣ Cube 3D
  const geometry = new THREE.BoxGeometry();
  const material = new THREE.MeshNormalMaterial();
  const cube = new THREE.Mesh(geometry, material);
  scene.add(cube);
  camera.position.z = 5;

  // 3️⃣ Animation
  function animate() {
    requestAnimationFrame(animate);
    cube.rotation.x += 0.01;
    cube.rotation.y += 0.01;
    renderer.render(scene, camera);
  }
  animate();

  // 4️⃣ Interaction souris
  document.addEventListener('mousemove', (e) => {
    cube.rotation.x = (e.clientY / window.innerHeight) * Math.PI;
    cube.rotation.y = (e.clientX / window.innerWidth) * Math.PI;
  });

  // 5️⃣ Redimensionnement
  window.addEventListener('resize', () => {
    camera.aspect = window.innerWidth/window.innerHeight;
    camera.updateProjectionMatrix();
    renderer.setSize(window.innerWidth, window.innerHeight);
  });
</script>
</body>
</html>
""", height=500)

